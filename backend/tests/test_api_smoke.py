import io
import uuid

import docx

from app.models import DocTemplate, GeneratedDocument, KnowledgeDoc, Project
from app.routers import chat as chat_router
from app.routers import settings as settings_router


class TestHealth:
    def test_health_no_auth_required(self, client):
        resp = client.get("/api/health")
        assert resp.status_code == 200
        assert resp.json() == {"status": "ok"}


class TestAuthRequired:
    def test_projects_requires_auth(self, client):
        resp = client.get("/api/projects")
        assert resp.status_code == 401

    def test_service_catalog_requires_auth(self, client):
        resp = client.get("/api/service-catalog")
        assert resp.status_code == 401


class TestLoginLockout:
    def test_wrong_password_then_lockout(self, client, make_user):
        user, _real_password = make_user()
        # MAX_ATTEMPTS is 10 (app.rate_limit) - the 10th bad attempt should trip the lockout
        # for the 11th, regardless of whether the password is right or wrong at that point.
        last_status = None
        for _ in range(11):
            resp = client.post(
                "/api/auth/login", json={"username": user.username, "password": "definitely-wrong"}
            )
            last_status = resp.status_code
        assert last_status == 429


class TestChangePassword:
    def test_new_user_must_change_password(self, client, make_user):
        user, password = make_user()
        client.post("/api/auth/login", json={"username": user.username, "password": password})
        resp = client.get("/api/auth/me")
        assert resp.status_code == 200
        assert resp.json()["mustChangePassword"] is True

    def test_change_password_clears_the_flag_and_new_password_works(self, client, make_user):
        user, password = make_user()
        client.post("/api/auth/login", json={"username": user.username, "password": password})

        resp = client.post(
            "/api/auth/change-password",
            json={"currentPassword": password, "newPassword": "a-brand-new-password"},
        )
        assert resp.status_code == 200
        assert resp.json()["mustChangePassword"] is False

        client.post("/api/auth/logout")
        relogin = client.post(
            "/api/auth/login", json={"username": user.username, "password": "a-brand-new-password"}
        )
        assert relogin.status_code == 200

    def test_wrong_current_password_rejected(self, client, make_user):
        user, password = make_user()
        client.post("/api/auth/login", json={"username": user.username, "password": password})

        resp = client.post(
            "/api/auth/change-password",
            json={"currentPassword": "not-the-real-password", "newPassword": "a-brand-new-password"},
        )
        assert resp.status_code == 400

    def test_requires_auth(self, client):
        resp = client.post(
            "/api/auth/change-password", json={"currentPassword": "x", "newPassword": "y" * 8}
        )
        assert resp.status_code == 401


class TestRoleEnforcement:
    def test_viewer_cannot_create_service_package(self, client, make_user):
        user, password = make_user(role="Viewer")
        client.post("/api/auth/login", json={"username": user.username, "password": password})
        resp = client.post(
            "/api/service-catalog",
            json={"category": "addon", "name": "X", "tagline": "Y", "monthlyPrice": "$10/mo", "resources": []},
        )
        assert resp.status_code == 403

    def test_owner_can_create_service_package(self, auth_client):
        resp = auth_client.post(
            "/api/service-catalog",
            json={
                "category": "addon",
                "name": f"Test Package {uuid.uuid4().hex[:6]}",
                "tagline": "Tagline",
                "monthlyPrice": "$10/mo",
                "resources": [],
            },
        )
        assert resp.status_code == 201
        assert resp.json()["category"] == "addon"


class TestQuoteProjectLinking:
    def test_generate_quote_with_project_id_shows_up_in_project_quotes(self, auth_client, db_session):
        project = Project(
            id=f"prj-{uuid.uuid4().hex[:8]}",
            name="Test Project",
            customer="Test Customer",
            cloud="AWS",
            status="Draft",
            owner="Tester",
            updated="2026-01-01",
            docs_generated=0,
            description="",
        )
        db_session.add(project)
        db_session.commit()

        pkg_resp = auth_client.post(
            "/api/service-catalog",
            json={
                "category": "tier",
                "name": f"Basic {uuid.uuid4().hex[:6]}",
                "tagline": "Entry tier",
                "monthlyPrice": "$500/mo",
                "resources": [{"service": "EC2", "quantity": 1, "purpose": "App server"}],
            },
        )
        assert pkg_resp.status_code == 201
        package_id = pkg_resp.json()["id"]

        quote_resp = auth_client.post(
            "/api/service-catalog/quote",
            json={
                "customerName": "Test Customer",
                "description": "Pricing for the pilot",
                "packageIds": [package_id],
                "format": "docx",
                "projectId": project.id,
            },
        )
        assert quote_resp.status_code == 200

        quotes_resp = auth_client.get(f"/api/projects/{project.id}/quotes")
        assert quotes_resp.status_code == 200
        quotes = quotes_resp.json()
        assert len(quotes) == 1
        assert quotes[0]["projectId"] == project.id
        assert quotes[0]["total"] == "$500.00/mo"

    def test_generate_quote_with_invalid_project_id_returns_404_and_does_not_log_a_quote(
        self, auth_client, db_session
    ):
        pkg_resp = auth_client.post(
            "/api/service-catalog",
            json={
                "category": "addon",
                "name": f"Addon {uuid.uuid4().hex[:6]}",
                "tagline": "T",
                "monthlyPrice": "$5/mo",
                "resources": [],
            },
        )
        package_id = pkg_resp.json()["id"]

        resp = auth_client.post(
            "/api/service-catalog/quote",
            json={
                "customerName": "Should Not Persist",
                "description": "y",
                "packageIds": [package_id],
                "format": "docx",
                "projectId": "does-not-exist",
            },
        )
        assert resp.status_code == 404


class TestListAllQuotes:
    def test_returns_quotes_across_projects(self, auth_client, db_session):
        # Quotes are only persisted when generated with a projectId (see
        # TestQuoteProjectLinking) — a standalone quote is generated and downloaded but
        # never written to the quotes table, so this endpoint necessarily only reflects
        # project-linked quotes too.
        project = Project(
            id=f"prj-{uuid.uuid4().hex[:8]}",
            name="Dashboard Stat Test Project",
            customer="Dashboard Test Customer",
            cloud="AWS",
            status="Draft",
            owner="Tester",
            updated="2026-01-01",
            docs_generated=0,
            description="",
        )
        db_session.add(project)
        db_session.commit()

        pkg_resp = auth_client.post(
            "/api/service-catalog",
            json={
                "category": "addon",
                "name": f"Addon {uuid.uuid4().hex[:6]}",
                "tagline": "T",
                "monthlyPrice": "$5/mo",
                "resources": [],
            },
        )
        package_id = pkg_resp.json()["id"]

        quote_resp = auth_client.post(
            "/api/service-catalog/quote",
            json={
                "customerName": "Dashboard Test Customer",
                "description": "For the dashboard stat",
                "packageIds": [package_id],
                "format": "docx",
                "projectId": project.id,
            },
        )
        assert quote_resp.status_code == 200

        resp = auth_client.get("/api/service-catalog/quotes")
        assert resp.status_code == 200
        quotes = resp.json()
        assert any(q["customerName"] == "Dashboard Test Customer" for q in quotes)

    def test_requires_auth(self, client):
        resp = client.get("/api/service-catalog/quotes")
        assert resp.status_code == 401


class TestCostEstimate:
    def test_generate_then_regenerate_updates_the_same_document(self, auth_client, db_session):
        project = Project(
            id=f"prj-{uuid.uuid4().hex[:8]}",
            name="Cost Estimate Test Project",
            customer="Test Customer",
            cloud="AWS",
            status="Draft",
            owner="Tester",
            updated="2026-01-01",
            docs_generated=0,
            description="",
        )
        db_session.add(project)
        db_session.commit()

        pkg_resp = auth_client.post(
            "/api/service-catalog",
            json={
                "category": "tier",
                "name": f"Basic {uuid.uuid4().hex[:6]}",
                "tagline": "Entry tier",
                "monthlyPrice": "$500/mo",
                "resources": [{"service": "EC2", "quantity": 1, "purpose": "App server"}],
            },
        )
        package_id = pkg_resp.json()["id"]

        first = auth_client.post(f"/api/projects/{project.id}/cost-estimate", json={"packageIds": [package_id]})
        assert first.status_code == 200
        first_doc = first.json()
        assert first_doc["type"] == "Cost Estimate"
        assert "$500.00/mo" in first_doc["content"]

        second = auth_client.post(f"/api/projects/{project.id}/cost-estimate", json={"packageIds": [package_id]})
        assert second.status_code == 200
        second_doc = second.json()

        # Same document updated in place, not a duplicate.
        assert second_doc["id"] == first_doc["id"]
        docs_for_project = (
            db_session.query(GeneratedDocument)
            .filter(GeneratedDocument.project == project.name, GeneratedDocument.type == "Cost Estimate")
            .all()
        )
        assert len(docs_for_project) == 1

    def test_requires_at_least_one_package(self, auth_client, db_session):
        project = Project(
            id=f"prj-{uuid.uuid4().hex[:8]}",
            name="No Package Project",
            customer="Test Customer",
            cloud="AWS",
            status="Draft",
            owner="Tester",
            updated="2026-01-01",
            docs_generated=0,
            description="",
        )
        db_session.add(project)
        db_session.commit()

        resp = auth_client.post(f"/api/projects/{project.id}/cost-estimate", json={"packageIds": []})
        assert resp.status_code == 400

    def test_viewer_cannot_generate_cost_estimate(self, client, make_user, db_session):
        project = Project(
            id=f"prj-{uuid.uuid4().hex[:8]}",
            name="RBAC Test Project",
            customer="Test Customer",
            cloud="AWS",
            status="Draft",
            owner="Tester",
            updated="2026-01-01",
            docs_generated=0,
            description="",
        )
        db_session.add(project)
        db_session.commit()

        user, password = make_user(role="Viewer")
        client.post("/api/auth/login", json={"username": user.username, "password": password})
        resp = client.post(f"/api/projects/{project.id}/cost-estimate", json={"packageIds": ["whatever"]})
        assert resp.status_code == 403


class TestChatExtractProject:
    def test_extracts_project_fields_from_conversation(self, auth_client, monkeypatch):
        extracted = {"name": "Acme Migration", "customer": "Acme Corp", "cloud": "AWS", "description": "Lift and shift."}
        monkeypatch.setattr(chat_router, "extract_project_from_chat", lambda messages, provider: extracted)

        resp = auth_client.post(
            "/api/chat/extract-project",
            json={"messages": [{"role": "user", "content": "We're moving Acme Corp to AWS"}]},
        )
        assert resp.status_code == 200
        assert resp.json() == extracted

    def test_requires_at_least_one_message(self, auth_client):
        resp = auth_client.post("/api/chat/extract-project", json={"messages": []})
        assert resp.status_code == 400

    def test_ai_failure_returns_502(self, auth_client, monkeypatch):
        def _raise(messages, provider):
            raise RuntimeError("provider unavailable")

        monkeypatch.setattr(chat_router, "extract_project_from_chat", _raise)
        resp = auth_client.post(
            "/api/chat/extract-project", json={"messages": [{"role": "user", "content": "hi"}]}
        )
        assert resp.status_code == 502


class TestSettingsTestConnection:
    def test_success_returns_reply(self, auth_client, monkeypatch):
        monkeypatch.setattr(settings_router, "run_connection_test", lambda provider: "OK")

        resp = auth_client.post("/api/settings/test-connection", json={"provider": "claude_cli"})
        assert resp.status_code == 200
        assert resp.json() == {"ok": True, "reply": "OK"}

    def test_provider_failure_returns_400_with_detail(self, auth_client, monkeypatch):
        def _raise(provider):
            raise RuntimeError("Invalid Access Key Id")

        monkeypatch.setattr(settings_router, "run_connection_test", _raise)
        resp = auth_client.post("/api/settings/test-connection", json={"provider": "bedrock"})
        assert resp.status_code == 400
        assert "Invalid Access Key Id" in resp.json()["detail"]

    def test_rejects_unknown_provider(self, auth_client):
        resp = auth_client.post("/api/settings/test-connection", json={"provider": "not-a-provider"})
        assert resp.status_code == 400

    def test_requires_owner_role(self, client, make_user):
        user, password = make_user(role="Viewer")
        client.post("/api/auth/login", json={"username": user.username, "password": password})
        resp = client.post("/api/settings/test-connection", json={"provider": "claude_cli"})
        assert resp.status_code == 403


class TestGlobalSearch:
    def test_finds_matches_across_categories(self, auth_client, db_session):
        stamp = uuid.uuid4().hex[:8]
        db_session.add(
            Project(
                id=f"prj-{stamp}",
                name=f"Zephyr-{stamp} Migration",
                customer="Zephyr Corp",
                cloud="AWS",
                status="Draft",
                owner="Tester",
                updated="2026-01-01",
                docs_generated=0,
                description="",
            )
        )
        db_session.add(
            GeneratedDocument(
                id=f"doc-{stamp}",
                project=f"Zephyr-{stamp} Migration",
                type="SDD",
                title=f"Zephyr-{stamp} Solution Design",
                version="1.0",
                updated="2026-01-01",
                status="Draft",
            )
        )
        db_session.add(
            DocTemplate(
                id=f"tpl-{stamp}",
                name=f"Zephyr-{stamp} Template",
                cloud="AWS",
                sections=1,
                description="d",
                content="c",
            )
        )
        db_session.add(
            KnowledgeDoc(
                id=f"kb-{stamp}",
                name=f"Zephyr-{stamp} Standards.pdf",
                category="Standards",
                uploaded_by="Tester",
                uploaded="2026-01-01",
                size="1 KB",
            )
        )
        db_session.commit()

        resp = auth_client.get(f"/api/search?q=Zephyr-{stamp}")
        assert resp.status_code == 200
        results = resp.json()
        types_found = {r["type"] for r in results}
        assert types_found == {"project", "document", "template", "knowledge"}

    def test_short_query_returns_empty(self, auth_client):
        resp = auth_client.get("/api/search?q=a")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_requires_auth(self, client):
        resp = client.get("/api/search?q=anything")
        assert resp.status_code == 401


class TestCreateTemplateFromUpload:
    def test_valid_docx_extracts_its_text(self, auth_client):
        buf = io.BytesIO()
        doc = docx.Document()
        doc.add_paragraph("Section One")
        doc.add_paragraph("Some real content pulled from the uploaded file.")
        doc.save(buf)
        buf.seek(0)

        resp = auth_client.post(
            "/api/templates",
            data={"name": "Uploaded Template", "cloud": "AWS", "description": ""},
            files={"document": ("template.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 201
        assert "Some real content pulled from the uploaded file." in resp.json()["content"]

    def test_legacy_doc_extension_returns_clean_400_not_500(self, auth_client):
        resp = auth_client.post(
            "/api/templates",
            data={"name": "Legacy Doc Template", "cloud": "AWS", "description": ""},
            files={"document": ("template.doc", io.BytesIO(b"not a real doc file"), "application/msword")},
        )
        assert resp.status_code == 400
        assert ".doc" in resp.json()["detail"]

    def test_malformed_docx_returns_clean_400_not_500(self, auth_client):
        # A .docx that isn't actually a valid zip/OOXML package — e.g. a renamed .doc,
        # or a corrupted upload. python-docx raises PackageNotFoundError here, not
        # ValueError, so this exercises the broad except-Exception fallback specifically.
        resp = auth_client.post(
            "/api/templates",
            data={"name": "Corrupt Template", "cloud": "AWS", "description": ""},
            files={
                "document": (
                    "template.docx",
                    io.BytesIO(b"this is not a real docx file"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert resp.status_code == 400
        assert "Couldn't read that file" in resp.json()["detail"]

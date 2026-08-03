import re
from datetime import date
from io import BytesIO
from typing import Callable, TypedDict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

PRICE_PATTERN = re.compile(r"[\d.]+")


def parse_monthly_price(price: str) -> float:
    """Extract the numeric value from a price string like '$1,200/mo' -> 1200.0."""
    digits = PRICE_PATTERN.findall(price.replace(",", ""))
    return float(digits[0]) if digits else 0.0


class QuotePackage(TypedDict):
    id: str
    name: str
    tagline: str
    monthlyPrice: str
    resources: list[dict]


# Mirrors lib/serviceCatalogTheme.ts on the frontend, so the branded proposal
# PDF uses the same per-package accent colors as the Service Catalog cards.
PACKAGE_ACCENTS = {
    "basic": "#0891b2",
    "intermediate": "#4f46e5",
    "advanced": "#7c3aed",
    "eks-service": "#326ce5",
    "ecs-service": "#ff9900",
    "managed-database-service": "#059669",
    "backup-dr-service": "#b45309",
    "security-compliance-service": "#be123c",
}
DEFAULT_ACCENT = "#4f46e5"


def build_quote_docx(customer_name: str, description: str, packages: list[QuotePackage]) -> bytes:
    """Build a customer-facing quote: cover info, one section per selected
    package with its resource bundle, and a total monthly price."""
    document = Document()
    document.add_heading("Service Quote", level=0)

    p = document.add_paragraph()
    p.add_run("Prepared for: ").bold = True
    p.add_run(customer_name or "Prospective Customer")
    p = document.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(date.today().isoformat())
    if description:
        p = document.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(description)

    document.add_paragraph()

    total = 0.0
    for pkg in packages:
        document.add_heading(pkg["name"], level=1)
        price_p = document.add_paragraph()
        price_p.add_run(pkg["monthlyPrice"]).bold = True
        total += parse_monthly_price(pkg["monthlyPrice"])

        if pkg.get("tagline"):
            tagline_p = document.add_paragraph()
            tagline_p.add_run(pkg["tagline"]).italic = True

        resources = pkg.get("resources") or []
        if resources:
            table = document.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            header_cells = table.rows[0].cells
            for idx, label in enumerate(["Resource", "Qty", "Purpose"]):
                header_cells[idx].text = label
            for r in resources:
                cells = table.add_row().cells
                cells[0].text = str(r.get("service", ""))
                cells[1].text = str(r.get("quantity", ""))
                cells[2].text = str(r.get("purpose", ""))

        document.add_paragraph()

    document.add_heading("Total Estimated Monthly Cost", level=1)
    total_p = document.add_paragraph()
    total_run = total_p.add_run(f"${total:,.2f}/mo")
    total_run.bold = True
    total_run.font.size = Pt(16)

    document.add_paragraph()
    footer = document.add_paragraph()
    footer.add_run(
        "This is an estimate based on the selected service packages. Final pricing may vary based on "
        "actual usage, region, and contract terms."
    ).italic = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _esc(text: str) -> str:
    """Escape text before embedding it in a reportlab Paragraph's mini-XML markup."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_quote_pdf(customer_name: str, description: str, packages: list[QuotePackage]) -> bytes:
    """Same content as build_quote_docx, rendered as a PDF via reportlab."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    total_style = ParagraphStyle("TotalPrice", parent=styles["Normal"], fontSize=16, spaceBefore=4)
    footer_style = ParagraphStyle("Footer", parent=styles["Italic"], fontSize=9, textColor=colors.HexColor("#64748b"))

    story = [
        Paragraph("Service Quote", styles["Title"]),
        Paragraph(f"<b>Prepared for:</b> {_esc(customer_name or 'Prospective Customer')}", styles["Normal"]),
        Paragraph(f"<b>Date:</b> {date.today().isoformat()}", styles["Normal"]),
    ]
    if description:
        story.append(Paragraph(f"<b>Description:</b> {_esc(description)}", styles["Normal"]))
    story.append(Spacer(1, 0.3 * inch))

    total = 0.0
    for pkg in packages:
        story.append(Paragraph(_esc(pkg["name"]), styles["Heading1"]))
        story.append(Paragraph(f"<b>{_esc(pkg['monthlyPrice'])}</b>", styles["Normal"]))
        total += parse_monthly_price(pkg["monthlyPrice"])

        if pkg.get("tagline"):
            story.append(Paragraph(f"<i>{_esc(pkg['tagline'])}</i>", styles["Normal"]))

        resources = pkg.get("resources") or []
        if resources:
            data = [["Resource", "Qty", "Purpose"]] + [
                [str(r.get("service", "")), str(r.get("quantity", "")), str(r.get("purpose", ""))] for r in resources
            ]
            table = Table(data, colWidths=[2.3 * inch, 0.5 * inch, 3.2 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(Spacer(1, 0.08 * inch))
            story.append(table)

        story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph("Total Estimated Monthly Cost", styles["Heading1"]))
    story.append(Paragraph(f"<b>${total:,.2f}/mo</b>", total_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        Paragraph(
            "This is an estimate based on the selected service packages. Final pricing may vary based on "
            "actual usage, region, and contract terms.",
            footer_style,
        )
    )

    doc.build(story)
    return buffer.getvalue()


def build_quote_proposal_pdf(customer_name: str, description: str, packages: list[QuotePackage]) -> bytes:
    """Branded letterhead-style proposal PDF — same content as build_quote_pdf,
    but with a company letterhead, serif headings, and per-package colored
    accent bars matching the Service Catalog's own card colors."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
    )
    styles = getSampleStyleSheet()
    brand_style = ParagraphStyle("Brand", parent=styles["Normal"], fontName="Times-Bold", fontSize=15)
    tagline_style = ParagraphStyle("BrandTagline", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#64748b"))
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#64748b"), alignment=2)
    eyebrow_style = ParagraphStyle(
        "Eyebrow", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#4f46e5"), spaceBefore=18
    )
    doc_title_style = ParagraphStyle(
        "DocTitle", parent=styles["Normal"], fontName="Times-Bold", fontSize=22, leading=27, spaceAfter=10
    )
    pkg_name_style = ParagraphStyle("PkgName", parent=styles["Normal"], fontName="Times-Bold", fontSize=14)
    pkg_price_style = ParagraphStyle("PkgPrice", parent=styles["Normal"], fontSize=13, alignment=2)
    pkg_tagline_style = ParagraphStyle("PkgTagline", parent=styles["Italic"], fontSize=9.5, textColor=colors.HexColor("#64748b"), spaceAfter=6)
    total_label_style = ParagraphStyle("TotalLabel", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#64748b"))
    total_value_style = ParagraphStyle("TotalValue", parent=styles["Normal"], fontName="Times-Bold", fontSize=26, textColor=colors.HexColor("#4f46e5"), alignment=2)
    footer_style = ParagraphStyle("Footer", parent=styles["Italic"], fontSize=8.5, textColor=colors.HexColor("#64748b"))

    def rule(hex_color: str = "#d8dee8", thickness: float = 0.75) -> Table:
        t = Table([[""]], colWidths=[6.4 * inch], rowHeights=[thickness])
        t.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), thickness, colors.HexColor(hex_color))]))
        return t

    # Letterhead
    letterhead = Table(
        [[
            [Paragraph("Meridian Cloud Services", brand_style), Paragraph("Managed Cloud Infrastructure", tagline_style)],
            [Paragraph(f"Date: {date.today().isoformat()}", meta_style)],
        ]],
        colWidths=[4.2 * inch, 2.2 * inch],
    )
    letterhead.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))

    story: list = [letterhead, Spacer(1, 6), rule(), Spacer(1, 4)]
    story.append(Paragraph("SERVICE QUOTE", eyebrow_style))
    story.append(Paragraph(_esc(f"A proposal for {customer_name or 'Prospective Customer'}"), doc_title_style))

    prepared_table_rows = [[Paragraph("<b>Prepared for</b>", meta_style), Paragraph(_esc(customer_name or "Prospective Customer"), styles["Normal"])]]
    if description:
        prepared_table_rows.append([Paragraph("<b>Description</b>", meta_style), Paragraph(_esc(description), styles["Normal"])])
    prepared_table = Table(prepared_table_rows, colWidths=[1.1 * inch, 5.3 * inch])
    prepared_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(prepared_table)
    story.append(Spacer(1, 0.25 * inch))

    total = 0.0
    for pkg in packages:
        accent = PACKAGE_ACCENTS.get(pkg.get("id", ""), DEFAULT_ACCENT)
        total += parse_monthly_price(pkg["monthlyPrice"])

        head = Table(
            [[Paragraph(_esc(pkg["name"]), pkg_name_style), Paragraph(_esc(pkg["monthlyPrice"]), pkg_price_style)]],
            colWidths=[4.6 * inch, 1.4 * inch],
        )
        head.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "BOTTOM")]))

        inner: list = [head]
        if pkg.get("tagline"):
            inner.append(Paragraph(_esc(pkg["tagline"]), pkg_tagline_style))

        resources = pkg.get("resources") or []
        if resources:
            cell_style = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=8.5)
            header_style = ParagraphStyle("CellHeader", parent=cell_style, fontName="Helvetica-Bold")
            data = [[Paragraph(h, header_style) for h in ("Resource", "Qty", "Purpose")]] + [
                [
                    Paragraph(_bold_to_reportlab(str(r.get("service", ""))), cell_style),
                    Paragraph(_bold_to_reportlab(str(r.get("quantity", ""))), cell_style),
                    Paragraph(_bold_to_reportlab(str(r.get("purpose", ""))), cell_style),
                ]
                for r in resources
            ]
            res_table = Table(data, colWidths=[2.1 * inch, 0.45 * inch, 3.05 * inch])
            res_table.setStyle(
                TableStyle(
                    [
                        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#64748b")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("LINEBELOW", (0, 0), (-1, 0), 0.75, colors.HexColor("#d8dee8")),
                        ("LINEBELOW", (0, 1), (-1, -2), 0.4, colors.HexColor("#e2e8f0")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            inner.append(res_table)

        # A narrow colored cell simulates the left accent bar from the web design.
        section = Table([[  "", inner]], colWidths=[0.06 * inch, 6.14 * inch])
        section.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(accent)),
                    ("LEFTPADDING", (1, 0), (1, 0), 14),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(section)
        story.append(Spacer(1, 0.22 * inch))

    story.append(rule("#1e2430", 1.25))
    story.append(Spacer(1, 6))
    total_row = Table(
        [[Paragraph("TOTAL ESTIMATED MONTHLY COST", total_label_style), Paragraph(f"${total:,.2f}/mo", total_value_style)]],
        colWidths=[3.5 * inch, 2.9 * inch],
    )
    total_row.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(total_row)
    story.append(Spacer(1, 0.35 * inch))
    story.append(
        Paragraph(
            "This is an estimate based on the selected service packages. Final pricing may vary based on "
            "actual usage, region, and contract terms.",
            footer_style,
        )
    )

    doc.build(story)
    return buffer.getvalue()


BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
HEADING_PATTERN = re.compile(r"^(#{1,4})\s+(.*)$")
CHECKBOX_PATTERN = re.compile(r"^- \[([ xX])\]\s+(.*)$")
IMAGE_PATTERN = re.compile(r"^!\[(.*?)\]\((.*?)\)$")


def _add_inline_runs(paragraph, text: str) -> None:
    """Split on **bold** markers and add runs, so inline emphasis survives the conversion."""
    pos = 0
    for m in BOLD_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(
    title: str, markdown_text: str, resolve_image: Callable[[str], bytes | None] | None = None
) -> bytes:
    """Convert our stored Markdown template content into a real .docx.

    Headings map to Word's built-in Heading 1-4 styles (not just bold text) —
    that's what makes Word's Navigation Pane and Table of Contents work, and
    is the whole point of this conversion instead of dumping raw Markdown.
    """
    document = Document()
    document.add_heading(title, level=0)

    lines = markdown_text.splitlines()
    in_code_block = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        rows = [r for r in table_rows if not re.fullmatch(r"[\s:|-]+", "|".join(r))]
        table_rows = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = document.add_table(rows=0, cols=cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for r in rows:
            cells = table.add_row().cells
            for idx in range(cols):
                cell_text = r[idx] if idx < len(r) else ""
                cells[idx].paragraphs[0].text = ""
                _add_inline_runs(cells[idx].paragraphs[0], cell_text)
        document.add_paragraph()

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = document.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_lines = []

    for raw_line in lines:
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
            else:
                flush_table()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if stripped.startswith("|"):
            table_rows.append([c.strip() for c in stripped.strip("|").split("|")])
            continue
        flush_table()

        if not stripped or stripped == "---":
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            document.add_heading(heading_match.group(2), level=level)
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            img_bytes = resolve_image(image_match.group(2)) if resolve_image else None
            if img_bytes:
                document.add_picture(BytesIO(img_bytes), width=Inches(6))
            else:
                p = document.add_paragraph()
                _add_inline_runs(p, stripped)
            continue

        checkbox_match = CHECKBOX_PATTERN.match(stripped)
        if checkbox_match:
            box = "☑" if checkbox_match.group(1).lower() == "x" else "☐"
            p = document.add_paragraph(style="List Bullet")
            _add_inline_runs(p, f"{box} {checkbox_match.group(2)}")
            continue

        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:])
            continue

        p = document.add_paragraph()
        _add_inline_runs(p, stripped)

    flush_table()
    flush_code()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _bold_to_reportlab(text: str) -> str:
    """Convert **bold** markers to reportlab's <b> mini-XML, escaping everything else."""
    parts = BOLD_PATTERN.split(text)
    out = []
    for i, part in enumerate(parts):
        escaped = _esc(part)
        out.append(f"<b>{escaped}</b>" if i % 2 == 1 else escaped)
    return "".join(out)


def markdown_to_pdf(
    title: str, markdown_text: str, resolve_image: Callable[[str], bytes | None] | None = None
) -> bytes:
    """Same parsing as markdown_to_docx, rendered as a PDF via reportlab."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    heading_styles = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16),
        2: ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13),
        3: ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11),
        4: ParagraphStyle("H4", parent=styles["Heading4"], fontSize=10),
    }
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], leftIndent=14, bulletIndent=0)
    code_style = ParagraphStyle("Code", parent=styles["Normal"], fontName="Courier", fontSize=8.5, leading=11)

    story: list = [Paragraph(_esc(title), styles["Title"]), Spacer(1, 0.2 * inch)]

    lines = markdown_text.splitlines()
    in_code_block = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        rows = [r for r in table_rows if not re.fullmatch(r"[\s:|-]+", "|".join(r))]
        table_rows = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table_cell_style = ParagraphStyle("TableCell", parent=styles["Normal"], fontSize=9)
        table_header_style = ParagraphStyle("TableHeader", parent=table_cell_style, fontName="Helvetica-Bold")
        data = [
            [
                Paragraph(_bold_to_reportlab(r[c]) if c < len(r) else "", table_header_style if i == 0 else table_cell_style)
                for c in range(cols)
            ]
            for i, r in enumerate(rows)
        ]
        t = Table(data, colWidths=[6.5 * inch / cols] * cols)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.15 * inch))

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        story.append(Paragraph(_esc("\n".join(code_lines)).replace("\n", "<br/>"), code_style))
        story.append(Spacer(1, 0.1 * inch))
        code_lines = []

    for raw_line in lines:
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
            else:
                flush_table()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if stripped.startswith("|"):
            table_rows.append([c.strip() for c in stripped.strip("|").split("|")])
            continue
        flush_table()

        if not stripped or stripped == "---":
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            story.append(Paragraph(_bold_to_reportlab(heading_match.group(2)), heading_styles[level]))
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            img_bytes = resolve_image(image_match.group(2)) if resolve_image else None
            if img_bytes:
                story.append(RLImage(BytesIO(img_bytes), width=6.5 * inch, height=8 * inch, kind="proportional"))
                story.append(Spacer(1, 0.15 * inch))
            else:
                story.append(Paragraph(_bold_to_reportlab(stripped), styles["Normal"]))
                story.append(Spacer(1, 0.05 * inch))
            continue

        checkbox_match = CHECKBOX_PATTERN.match(stripped)
        if checkbox_match:
            box = "☑" if checkbox_match.group(1).lower() == "x" else "☐"
            story.append(Paragraph(f"{box} {_bold_to_reportlab(checkbox_match.group(2))}", bullet_style))
            continue

        if stripped.startswith("- "):
            story.append(Paragraph(f"• {_bold_to_reportlab(stripped[2:])}", bullet_style))
            continue

        story.append(Paragraph(_bold_to_reportlab(stripped), styles["Normal"]))
        story.append(Spacer(1, 0.05 * inch))

    flush_table()
    flush_code()

    doc.build(story)
    return buffer.getvalue()
